import os
import re
import tempfile
from typing import Dict, Optional, List
from pathlib import Path
import docx
import striprtf
from striprtf.striprtf import rtf_to_text


class FileResumeParser:
    """Парсер для файлов резюме в форматах RTF и DOCX"""
    
    # Ключевые слова для поиска секций
    SECTION_KEYWORDS = {
        'experience': ['опыт работы', 'experience', 'трудовая деятельность', 'работал', 'работала'],
        'education': ['образование', 'education', 'университет', 'институт', 'колледж'],
        'skills': ['навыки', 'skills', 'технологии', 'technologies', 'компетенции'],
        'languages': ['языки', 'languages', 'английский', 'english'],
        'projects': ['проекты', 'projects', 'портфолио', 'portfolio']
    }
    
    @staticmethod
    def parse_rtf(file_path: str) -> str:
        """Парсинг RTF файла"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                rtf_content = file.read()
            return rtf_to_text(rtf_content)
        except Exception as e:
            print(f"Ошибка парсинга RTF: {e}")
            return ""
    
    @staticmethod
    def parse_docx(file_path: str) -> str:
        """Парсинг DOCX файла"""
        try:
            # Проверяем существование файла
            if not os.path.exists(file_path):
                print(f"Файл не найден: {file_path}")
                return ""
            
            # Проверяем размер файла
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                print(f"Файл пустой: {file_path}")
                return ""
            
            print(f"Парсинг DOCX файла: {file_path}, размер: {file_size} байт")
            
            # Открываем документ
            doc = docx.Document(file_path)
            
            # Извлекаем текст из параграфов
            text_parts = []
            
            # Обрабатываем параграфы
            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                if para_text:
                    text_parts.append(para_text)
                    print(f"Параграф {i}: {para_text[:50]}...")
            
            # Обрабатываем таблицы
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        text_parts.append(" | ".join(row_texts))
            
            # Обрабатываем заголовки
            for section in doc.sections:
                header = section.header
                if header:
                    for paragraph in header.paragraphs:
                        header_text = paragraph.text.strip()
                        if header_text:
                            text_parts.append(header_text)
            
            result = '\n'.join(text_parts)
            print(f"Извлечено {len(result)} символов текста")
            return result
            
        except Exception as e:
            print(f"Ошибка парсинга DOCX: {e}")
            import traceback
            traceback.print_exc()
            return ""
    
    @staticmethod
    def find_sections(text: str) -> Dict[str, str]:
        """Поиск секций в резюме"""
        sections = {}
        lines = text.split('\n')
        
        current_section = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Проверяем, является ли строка заголовком секции
            is_section_header = False
            for section_name, keywords in FileResumeParser.SECTION_KEYWORDS.items():
                for keyword in keywords:
                    if re.search(rf'\b{re.escape(keyword)}\b', line, re.IGNORECASE):
                        # Сохраняем предыдущую секцию
                        if current_section and current_content:
                            sections[current_section] = '\n'.join(current_content)
                        
                        current_section = section_name
                        current_content = []
                        is_section_header = True
                        break
                if is_section_header:
                    break
            
            if not is_section_header and current_section:
                current_content.append(line)
        
        # Сохраняем последнюю секцию
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    @staticmethod
    def extract_title(text: str) -> str:
        """Извлечение заголовка резюме"""
        lines = text.split('\n')
        
        # Ищем заголовок в первых строках
        for line in lines[:10]:
            line = line.strip()
            if line and len(line) < 100:
                # Исключаем номера страниц, даты и служебную информацию
                if not re.match(r'^\d+$', line) and not re.match(r'^\d{1,2}[./]\d{1,2}[./]\d{2,4}$', line):
                    # Проверяем, что это не служебная информация
                    if not any(keyword in line.lower() for keyword in ['резюме', 'resume', 'страница', 'page']):
                        return line
        
        return "Кандидат"
    
    @staticmethod
    def extract_skills(text: str, sections: Dict[str, str]) -> str:
        """Извлечение навыков"""
        # Сначала ищем в секции skills
        if 'skills' in sections:
            skills_text = sections['skills']
            # Очищаем от лишних символов
            skills_text = re.sub(r'[•\-\*]', ',', skills_text)
            skills_text = re.sub(r'\s+', ' ', skills_text)
            return skills_text.strip()
        
        # Если секции нет, ищем ключевые слова в тексте
        tech_keywords = [
            'Python', 'Java', 'JavaScript', 'TypeScript', 'C++', 'C#', '.NET',
            'SQL', 'PostgreSQL', 'MySQL', 'MongoDB', 'Redis', 'Elasticsearch',
            'Docker', 'Kubernetes', 'Jenkins', 'GitLab', 'GitHub', 'Git',
            'AWS', 'Azure', 'GCP', 'Linux', 'Windows', 'MacOS',
            'FastAPI', 'Django', 'Flask', 'Spring', 'React', 'Vue', 'Angular',
            'Node.js', 'Express', 'REST API', 'GraphQL', 'Microservices',
            'Excel', 'PowerPoint', 'Word', 'Visio', 'Jira', 'Confluence',
            'Scrum', 'Agile', 'Kanban', 'CI/CD', 'DevOps', 'SRE',
            'Machine Learning', 'Data Science', 'Pandas', 'NumPy', 'Sklearn',
            'TensorFlow', 'PyTorch', 'Tableau', 'Power BI', 'SAS', 'R',
            'Tableau', 'Looker', 'Metabase', 'Apache Spark', 'Hadoop',
            'Kafka', 'RabbitMQ', 'Nginx', 'Apache', 'Tomcat', 'WildFly',
            # Добавляем IT-навыки из вашего резюме
            'Active Directory', 'Windows Server', 'ITIL', 'Helpdesk', 'ITSM',
            'Service Desk', 'GPO', 'DHCP', 'DNS', 'VBScript', 'PowerShell',
            'HP Service Manager', 'Atlassian Jira', 'WDS', 'Mac OS'
        ]
        
        found_skills = []
        text_lower = text.lower()
        
        for keyword in tech_keywords:
            if re.search(rf'\b{re.escape(keyword)}\b', text, re.IGNORECASE):
                found_skills.append(keyword)
        
        return ', '.join(found_skills) if found_skills else ""
    
    @staticmethod
    def extract_experience(text: str, sections: Dict[str, str]) -> str:
        """Извлечение опыта работы"""
        # Сначала ищем в секции experience
        if 'experience' in sections:
            return sections['experience']
        
        # Ищем паттерны опыта в тексте
        experience_patterns = [
            r'Опыт работы[:\s]*(.*?)(?=\n\n|\n[A-ZА-Я]|$)',
            r'Experience[:\s]*(.*?)(?=\n\n|\n[A-ZА-Я]|$)',
            r'Работал[:\s]*(.*?)(?=\n\n|\n[A-ZА-Я]|$)',
            r'Работала[:\s]*(.*?)(?=\n\n|\n[A-ZА-Я]|$)',
            r'Трудовая деятельность[:\s]*(.*?)(?=\n\n|\n[A-ZА-Я]|$)',
        ]
        
        for pattern in experience_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                return match.group(1).strip()
        
        # Если не нашли по паттернам, ищем блок с опытом работы
        lines = text.split('\n')
        experience_lines = []
        in_experience_section = False
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Проверяем, начинается ли секция опыта
            if any(keyword in line.lower() for keyword in ['опыт работы', 'experience', 'трудовая деятельность']):
                in_experience_section = True
                continue
            
            # Если мы в секции опыта, собираем строки
            if in_experience_section:
                # Проверяем, не начинается ли новая секция
                if any(keyword in line.lower() for keyword in ['образование', 'education', 'навыки', 'skills']):
                    break
                experience_lines.append(line)
        
        if experience_lines:
            return '\n'.join(experience_lines)
        
        return ""
    
    @staticmethod
    def extract_resume_data(text: str) -> Dict[str, str]:
        """Извлечение структурированных данных из текста резюме"""
        print(f"Обработка текста длиной {len(text)} символов")
        
        # Очистка текста
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Поиск секций
        sections = FileResumeParser.find_sections(text)
        print(f"Найдено секций: {list(sections.keys())}")
        
        # Извлечение данных
        title = FileResumeParser.extract_title(text)
        experience = FileResumeParser.extract_experience(text, sections)
        skills = FileResumeParser.extract_skills(text, sections)
        
        print(f"Извлечено - Заголовок: {title}")
        print(f"Извлечено - Опыт: {len(experience)} символов")
        print(f"Извлечено - Навыки: {skills}")
        
        # Ограничиваем длину полей
        if experience and len(experience) > 500:
            experience = experience[:500] + "..."
        if skills and len(skills) > 300:
            skills = skills[:300] + "..."
        
        return {
            "title": title,
            "experience": experience,
            "skills": skills,
            "full_text": text[:1000],  # Первые 1000 символов для дополнительного анализа
            "sections": sections  # Сохраняем найденные секции
        }
    
    @staticmethod
    def parse_file(file_path: str) -> Optional[Dict[str, str]]:
        """Основной метод парсинга файла"""
        if not os.path.exists(file_path):
            print(f"Файл не найден: {file_path}")
            return None
        
        file_ext = Path(file_path).suffix.lower()
        print(f"Парсинг файла: {file_path}, расширение: {file_ext}")
        
        if file_ext == '.rtf':
            text = FileResumeParser.parse_rtf(file_path)
        elif file_ext == '.docx':
            text = FileResumeParser.parse_docx(file_path)
        else:
            print(f"Неподдерживаемый формат файла: {file_ext}")
            return None
        
        if not text:
            print("Не удалось извлечь текст из файла")
            return None
        
        print(f"Успешно извлечено {len(text)} символов текста")
        return FileResumeParser.extract_resume_data(text) 